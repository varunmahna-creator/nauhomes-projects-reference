"""Generate the ZuildUp dev-team handover document as a .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, sys

OUT = "/opt/openclaw/workspace/nauhomes-projects-export/Nauhomes_Projects_Module_Handover.docx"

doc = Document()

# ---- Styling helpers ---------------------------------------------------
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x1F, 0x44)  # ZuildUp navy
    return h

def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def add_code(text, caption=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    # subtle background via pPr shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

# ---- TITLE PAGE --------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Projects + Timeline Module")
tr.font.size = Pt(28); tr.bold = True
tr.font.color.rgb = RGBColor(0x0A, 0x1F, 0x44)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Reference Code Handover for ZuildUp Dev Team")
sr.font.size = Pt(14)
sr.font.color.rgb = RGBColor(0xC9, 0xA2, 0x4D)  # gold

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("\nPrepared by: Varun Mahna\nSource project: nauhomes.com\nDate: 4 May 2026\nSnapshot commit: 725ee70")
mr.font.size = Pt(11)

doc.add_page_break()

# ---- SECTION 1: WHAT THIS IS ------------------------------------------
add_heading("1. What this is", 1)
add_para(
    "This handover contains the complete source code for the Projects feature on "
    "nauhomes.com, including the construction-timeline editor used by the admin "
    "and the public-facing project pages."
)
add_para(
    "Goal: port this feature into ZuildUp's platform so site visitors can browse "
    "ongoing/completed projects and view live construction timelines, while the "
    "ZuildUp team manages everything from a single admin dashboard."
)
add_para("What's included:", bold=True)
for item in [
    "Postgres schema + CRUD layer (projects table with JSONB timeline)",
    "Public REST API (GET list, GET by slug)",
    "Admin REST API (POST / PUT / DELETE) protected by HTTP-only cookie auth",
    "Admin dashboard UI (1.7k LOC) — Next.js client component with tabs for Projects, Timeline, Settings, Leads, Media, Testimonials",
    "Public listing page + project detail page (with gallery, specs, timeline UI)",
    "Image upload pipeline — small (<4 MB) via API + large (>4 MB) via direct client-upload to Vercel Blob",
    "Auth middleware + login/logout endpoints",
    "Testimonials module (separate from projects, used by the home page)",
    "TypeScript type definitions for Project + TimelineEntry",
]:
    add_bullet(item)

# ---- SECTION 2: WHERE TO GET THE CODE ---------------------------------
add_heading("2. Where to get the code", 1)
add_para("All files are pushed to a dedicated branch on the nauhomes repository:")
add_code("https://github.com/varunmahna-creator/nauhomes/tree/projects-module-export",
         caption="GitHub branch")
add_para("Clone command:")
add_code("git clone -b projects-module-export \\\n  git@github.com:varunmahna-creator/nauhomes.git \\\n  zuildup-projects-reference",
         caption="terminal")
add_para(
    "The repo is private. Varun will add the dev-team GitHub usernames as "
    "collaborators with read-only access. Send your GitHub usernames to him.",
    italic=True,
)

# ---- SECTION 3: STACK & DEPENDENCIES ----------------------------------
add_heading("3. Stack & dependencies", 1)
for item in [
    "Next.js 15 (App Router) + TypeScript",
    "Vercel Postgres (@vercel/postgres) — tagged-template SQL, serverless-friendly",
    "Vercel Blob (@vercel/blob) — image/video storage with two-tier upload (server-proxied for small, client-direct for large)",
    "Tailwind CSS for styling",
    "Framer Motion for animations",
    "HTTP-only signed cookie for admin auth (no NextAuth dependency — simple and lean)",
]:
    add_bullet(item)
add_para(
    "If ZuildUp is on a different stack (e.g. Supabase + S3), the API contract "
    "and DB schema are platform-neutral. Swap @vercel/postgres for pg / "
    "postgres.js, and @vercel/blob for Supabase Storage / S3 / R2. The "
    "CRUD logic in lib/projects-db.ts stays nearly identical.",
    italic=True,
)

# ---- SECTION 4: ENV VARS ----------------------------------------------
add_heading("4. Environment variables required", 1)
add_code(
    "POSTGRES_URL=postgres://...                  # Postgres connection string\n"
    "BLOB_READ_WRITE_TOKEN=vercel_blob_rw_...     # Vercel Blob token\n"
    "ADMIN_PASSWORD=<choose-a-strong-password>    # for /admin login\n"
    "ADMIN_SESSION_SECRET=<random-32-char-string> # signs the admin cookie",
    caption=".env / .env.production"
)

# ---- SECTION 5: DATABASE SCHEMA ---------------------------------------
add_heading("5. Database schema", 1)
add_para(
    "The projects table auto-bootstraps on the first DB call via "
    "ensureSchema() in lib/projects-db.ts. No manual migration step needed."
)
add_code(
    "CREATE TABLE IF NOT EXISTS projects (\n"
    "  slug              TEXT PRIMARY KEY,\n"
    "  title             TEXT NOT NULL,\n"
    "  subtitle          TEXT DEFAULT '',\n"
    "  location          TEXT NOT NULL,           -- 'delhi' | 'bali'\n"
    "  location_label    TEXT DEFAULT '',         -- e.g. 'Greater Kailash, New Delhi'\n"
    "  status            TEXT DEFAULT 'ongoing',  -- 'ongoing' | 'completed'\n"
    "  type              TEXT DEFAULT '',         -- 'Independent Villa', etc.\n"
    "  area              TEXT DEFAULT '',         -- '8,500 sq ft'\n"
    "  year              TEXT DEFAULT '',\n"
    "  thumbnail         TEXT DEFAULT '',\n"
    "  gallery           JSONB DEFAULT '[]',      -- ProjectImage[]\n"
    "  floor_plans       JSONB DEFAULT '[]',\n"
    "  tour_embed_url    TEXT,                    -- Matterport / 3D tour URL\n"
    "  description       TEXT DEFAULT '',\n"
    "  highlights        JSONB DEFAULT '[]',      -- string[]\n"
    "  amenities         JSONB DEFAULT '[]',\n"
    "  specs             JSONB DEFAULT '{}',      -- key-value pairs\n"
    "  timeline          JSONB DEFAULT '[]',      -- TimelineEntry[]  <-- KEY\n"
    "  virtual_tour_videos JSONB DEFAULT '[]',\n"
    "  created_at        TIMESTAMPTZ DEFAULT NOW(),\n"
    "  updated_at        TIMESTAMPTZ DEFAULT NOW()\n"
    ");\n\n"
    "CREATE INDEX projects_location_idx ON projects(location);\n"
    "CREATE INDEX projects_status_idx   ON projects(status);",
    caption="SQL"
)

add_heading("TimelineEntry shape (stored in projects.timeline as JSONB)", 2)
add_code(
    "interface TimelineEntry {\n"
    "  date: string;          // free-form: 'Mar 2024', 'Phase 3 — Apr 2024'\n"
    "  title: string;         // 'Foundation laid', 'Roof slab cast'\n"
    "  description: string;\n"
    "  images: ProjectImage[]; // [{src: blobUrl, alt: '...'}]\n"
    "  videos?: ProjectImage[];\n"
    "}",
    caption="TypeScript"
)
add_para(
    "Storing the entire timeline array in JSONB keeps the schema simple and the "
    "admin UI dead easy — the whole timeline is replaced on save. If a project "
    "ever needs to fan out to thousands of phase entries, split into a "
    "project_timeline child table. We never hit that limit."
)

# ---- SECTION 6: FILE MAP ----------------------------------------------
add_heading("6. File map", 1)
add_para("All paths are relative to the export branch root.")

groups = [
    ("Data layer", [
        ("lib/projects-db.ts",        "Schema bootstrap + CRUD (single source of truth)"),
        ("lib/testimonials.ts",       "Testimonials data layer (separate Postgres table)"),
        ("types.ts",                  "Project + TimelineEntry + Testimonial interfaces"),
    ]),
    ("API routes", [
        ("api/projects/route.ts",         "GET (public list) / POST (admin create)"),
        ("api/projects-slug/route.ts",    "GET (public) / PUT / DELETE (admin)"),
        ("api/testimonials/route.ts",     "GET (public) / POST (admin)"),
        ("api/testimonials-id/route.ts",  "PUT / DELETE (admin)"),
        ("api/admin/login/route.ts",      "Password → HTTP-only cookie"),
        ("api/admin/logout/route.ts",     "Clears cookie"),
        ("api/upload/route.ts",           "Small image upload (<4 MB) → Vercel Blob"),
        ("api/upload-large/route.ts",     "Server-issued upload token for large files"),
        ("api/blob-upload/route.ts",      "Vercel Blob direct client-upload handler"),
        ("api/media/route.ts",            "List/manage uploaded blobs"),
        ("api/media/[id]/route.ts",       "Delete a blob"),
    ]),
    ("Auth & middleware", [
        ("middleware.ts", "Guards /admin/* and admin API routes via cookie check"),
    ]),
    ("Admin UI", [
        ("admin/page.tsx",          "Thin server wrapper (auth gate)"),
        ("admin/AdminDashboard.tsx","THE admin UI (~1,674 LOC). Has Projects tab with timeline editor."),
    ]),
    ("Public pages", [
        ("page-projects/page.tsx",          "/projects list page (server component)"),
        ("page-project-slug/page.tsx",      "/projects/[slug] detail page (server component)"),
        ("components/ProjectsPageClient.tsx",   "Interactive grid + location/status filters"),
        ("components/ProjectDetailClient.tsx",  "Gallery + timeline UI + specs"),
        ("components/FeaturedProjects.tsx",     "Homepage carousel"),
        ("components/ProcessTimeline.tsx",      "Homepage 6-step process (separate from per-project timeline)"),
        ("components/Testimonials.tsx",         "Homepage testimonials carousel"),
    ]),
]
for title_, rows in groups:
    add_heading(title_, 2)
    for path, desc in rows:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path); r.font.name = "Consolas"; r.bold = True; r.font.size = Pt(10)
        p.add_run(" — " + desc).font.size = Pt(10)

# ---- SECTION 7: ADMIN UI FLOW -----------------------------------------
add_heading("7. Admin UI flow (the part to demo)", 1)
steps = [
    "Admin logs in at /admin/login. POST /api/admin/login validates password and sets an HTTP-only cookie. The middleware (middleware.ts) rejects any non-cookie request to /admin/* and admin API routes.",
    "/admin renders AdminDashboard.tsx. Sidebar tabs: Projects, Settings, Leads, Media, Testimonials.",
    "Projects tab: list of all projects with inline DELETE; 'Add Project' / 'Edit' opens a form with sections — Basic info, Description/Highlights/Amenities/Specs, Thumbnail, Gallery, Floor plans, Virtual-tour videos / Matterport URL, and the Timeline editor.",
    "Timeline editor: 'Add phase' button creates a new row with date / title / description fields and a multi-image drop zone. Reorder via drag handles. Delete per row.",
    "Save → PUT /api/projects/:slug (or POST /api/projects for new). Server reads existing record, merges JSONB fields, writes back atomically.",
    "Public pages (/projects, /projects/[slug]) are server components that call getProjects() / getProjectBySlug() directly from lib/projects-db.ts. Set export const dynamic = 'force-dynamic' if you want edits to appear instantly without revalidation.",
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(s).font.size = Pt(11)

# ---- SECTION 8: IMAGE UPLOAD ARCHITECTURE -----------------------------
add_heading("8. Image upload architecture (CRITICAL)", 1)
add_para(
    "Vercel functions have a 4.5 MB request body cap. A single phase photo from "
    "a modern phone exceeds this. We learned this the hard way (413 errors in "
    "production for any timeline photo over 4 MB). The fix is two-tier:",
    bold=True,
)
add_para("Small files (<4 MB):", bold=True)
add_bullet("Browser sends multipart POST to /api/upload. Server forwards bytes to @vercel/blob.put(). Returns the public URL.")
add_para("Large files (>4 MB):", bold=True)
add_bullet("Browser requests a signed upload token from POST /api/upload-large.")
add_bullet("Browser uses @vercel/blob/client.upload() to send the file directly to Vercel Blob — server is bypassed entirely.")
add_bullet("Server's /api/blob-upload endpoint validates the upload via webhook on completion.")
add_para(
    "If ZuildUp uses S3/R2/Supabase Storage instead of Vercel Blob, replicate "
    "the same two-tier pattern with pre-signed PUT URLs. Do NOT route large "
    "uploads through your own API — every image >4 MB will 413."
)

# ---- SECTION 9: PORTING CHECKLIST -------------------------------------
add_heading("9. Porting checklist for ZuildUp", 1)
items = [
    "Provision a Postgres database (Vercel Postgres, Supabase Postgres, or self-hosted). Set POSTGRES_URL.",
    "If keeping Vercel Blob, set BLOB_READ_WRITE_TOKEN. Otherwise, swap api/upload* and api/blob-upload to use S3/R2/Supabase Storage with pre-signed URLs.",
    "If ZuildUp already has user auth (e.g. Supabase Auth), replace middleware.ts cookie check with a getUser() + role check (role IN ('admin','superadmin')). Drop api/admin/login and api/admin/logout.",
    "Drop AdminDashboard.tsx into the ZuildUp admin shell. The component is monolithic but self-contained — its only dependencies are the API routes and Tailwind classes.",
    "Run schema bootstrap by calling any read endpoint once (ensureSchema() runs lazily on first call), or run the SQL manually in Section 5.",
    "Update API base URLs in AdminDashboard.tsx if you mount it under a different prefix (e.g. /admin/projects instead of /admin).",
    "Add 'export const dynamic = \"force-dynamic\"' to /projects and /projects/[slug] if you need edits to appear instantly without manual revalidation.",
    "Test end-to-end: create a project from admin → upload a 6 MB phase photo → visit public URL → confirm photo displays. (See gotcha section.)",
]
for s in items:
    p = doc.add_paragraph(style="List Number")
    p.add_run(s).font.size = Pt(11)

# ---- SECTION 10: GOTCHAS / LESSONS ------------------------------------
add_heading("10. Gotchas (battle-tested — read this)", 1)

gotchas = [
    ("Vercel function body cap is 4.5 MB",
     "Any image >4 MB MUST go through the client-direct upload path (api/upload-large + @vercel/blob/client). Don't pipe through your own API route — every video upload will 413. (We hit this in production with the timeline photos. Hours lost.)"),
    ("Vercel deploy author allowlist",
     "If ZuildUp's Vercel project trusts only specific commit authors, agent-authored commits get silently 'Blocked'. Builds never run. Configure trusted authors or commit as the trusted email."),
    ("In-memory storage trap",
     "Don't use Map / module-scope variables for state on Vercel serverless. Each request can hit a different instance — saves on instance A, reads on instance B serve defaults. Use Postgres (or Redis) from day one. (We hit this with settings + leads + projects all three. Same fix pattern each time.)"),
    ("force-dynamic on DB-reading routes",
     "Without 'export const dynamic = \"force-dynamic\"' on routes that read from Postgres, Next.js may serve stale cached HTML. Apply to /contact, /api/settings, /api/leads, and optionally /projects/[slug]."),
    ("JSONB merging on UPDATE",
     "updateProject() does merged = { ...existing, ...partial } in JS, then writes the full record. This prevents partial updates from clobbering nested arrays like timeline or gallery. If you write a custom update path, preserve this pattern."),
    ("Hardcoded ranges are silent killers",
     "When iterating arrays / fetching paginated data, ask 'what happens if data exceeds my limit?' Silent truncation has burned us (Sales OS sheet ingest dropped 363 leads with no error). Use open-ended ranges or paginate properly."),
    ("Client hydration flash",
     "Client components reading /api/* via fetch will flash stale fallback content on first paint, visible to users and SEO crawlers. Pass server-fetched initial state as a prop. See useContactInfo hook pattern in nauhomes for reference."),
]
for title_, body in gotchas:
    p = doc.add_paragraph()
    r = p.add_run("• " + title_)
    r.bold = True; r.font.size = Pt(11)
    pp = doc.add_paragraph(body)
    pp.paragraph_format.left_indent = Inches(0.25)
    for run in pp.runs:
        run.font.size = Pt(10)

# ---- SECTION 11: DEPENDENCIES -----------------------------------------
add_heading("11. NPM dependencies (from nauhomes package.json)", 1)
add_para("Key runtime deps to install:")
add_code(
    "npm install \\\n"
    "  next@^15 \\\n"
    "  react@^19 react-dom@^19 \\\n"
    "  @vercel/postgres \\\n"
    "  @vercel/blob \\\n"
    "  framer-motion \\\n"
    "  lucide-react \\\n"
    "  tailwindcss",
    caption="terminal"
)
add_para("The exact pinned versions are in package.json on the export branch.")

# ---- SECTION 12: CONTACT / NEXT STEPS ---------------------------------
add_heading("12. Next steps & contact", 1)
add_bullet("Send GitHub usernames of dev-team members who need read access to nauhomes repo. Varun will add them as collaborators.")
add_bullet("Once they have access, clone the projects-module-export branch.")
add_bullet("Read this doc + the README.md on the branch in parallel — README has the same info but in markdown for quick reference in editors.")
add_bullet("For questions about implementation choices: ask Varun, who can loop in IRAAJ (the AI assistant that built the original nauhomes module) for architectural follow-ups.")

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run("— End of handover —")
fr.italic = True; fr.font.size = Pt(10)
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save(OUT)
print("OK", OUT, os.path.getsize(OUT), "bytes")
